import docx

document = docx.Document()

cars = ['Ford', 'Honda', 'Ferrari', 'Toyota']

document.add_paragraph('Cars', 'Heading 1')

# Add a paragraph for each list item
# Set any style you like
for car in cars:
    document.add_paragraph(car, 'Heading 3')

# Save
document.save('cars.docx')



