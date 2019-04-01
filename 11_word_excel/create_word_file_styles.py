import docx

document = docx.Document()

# Add a title
document.add_paragraph('Hello Word!', 'Title')

# Add a paragraph, specify style.
document.add_paragraph('By Clara', 'Heading 1')

document.add_paragraph('This is a word document created by Python, using the python-docx library.')

document.add_paragraph('Python programming is challenging!', 'Intense Quote')

document.add_paragraph('This paragraph has some ')
last_paragraph = document.paragraphs[-1]    # Get the last paragraph added
last_paragraph.add_run('bold, underlined text')         # Add a run to this paragraph
last_run = last_paragraph.runs[-1]          # Get the last run
last_run.bold = True                        # Set the bold attribute to True
last_run.underline = True                   # Set the underline attribute to True
last_paragraph.add_run(' in the middle. ')  # add another run

# Save
document.save('hello_word_style.docx')



