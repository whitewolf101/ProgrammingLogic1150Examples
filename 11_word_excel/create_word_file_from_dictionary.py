import docx

document = docx.Document()

electric_cars = {
    'Chevrolet': 'Volt',
    'Nissan': 'Leaf',
    'Tesla': 'Model S',
    'Volkswagen': 'e-Golf'
}

document.add_paragraph('Electric Cars', 'Heading 1')

# Add a paragraph for each list item
# Set any style you like
for make, model in electric_cars.items():
    document.add_paragraph(make, 'Heading 3')
    document.add_paragraph(f'An electric car made by {make} is {model}.')

# Save
document.save('electric_cars.docx')



