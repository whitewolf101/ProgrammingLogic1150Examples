import docx

document = docx.Document('IT_Sample_Resume.docx')

print(document.paragraphs)

for para in document.paragraphs:
    print(para.text)
